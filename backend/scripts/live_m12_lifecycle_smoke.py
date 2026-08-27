import asyncio
import json
import os
import re
import subprocess
import sys
import tempfile
import threading
import time
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from uuid import UUID, uuid4

import httpx
import jwt
from cryptography.hazmat.primitives.asymmetric import rsa
from docx import Document as DocxDocument
from sqlalchemy import func, select

ISSUER = "http://host.docker.internal:8765"
AUDIENCE = "production-rag-assistant-api"
API = "http://127.0.0.1:18000"
PROJECT = "m12-lifecycle-smoke"
POSTGRES_PORT = "15433"
TOKEN_LIFETIME_SECONDS = 300
REPOSITORY = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPOSITORY / "backend"))


def public_jwk(private_key: object) -> dict[str, object]:
    value = json.loads(jwt.algorithms.RSAAlgorithm.to_jwk(private_key.public_key()))
    value.update({"kid": "m12-key", "alg": "RS256", "use": "sig"})
    return value


def access_token(private_key: object, subject: str) -> str:
    now = int(time.time())
    return jwt.encode(
        {
            "iss": ISSUER,
            "sub": subject,
            "aud": AUDIENCE,
            "iat": now,
            "nbf": now - 1,
            "exp": now + TOKEN_LIFETIME_SECONDS,
            "scope": "rag:access",
        },
        private_key,
        algorithm="RS256",
        headers={"kid": "m12-key", "typ": "at+jwt"},
    )


class QuietJWKSHandler(BaseHTTPRequestHandler):
    jwks: dict[str, object] = {"keys": []}

    def do_GET(self) -> None:  # noqa: N802
        if self.path != "/jwks":
            self.send_response(404)
            self.end_headers()
            return
        body = json.dumps(self.jwks).encode()
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format: str, *args: object) -> None:
        return None


def environment() -> dict[str, str]:
    return {
        **os.environ,
        "COMPOSE_PROJECT_NAME": PROJECT,
        "POSTGRES_PORT": POSTGRES_PORT,
        "API_PORT": "18000",
        "DATABASE_URL": (
            "postgresql+asyncpg://rag_assistant_dev:rag_assistant_dev@"
            "postgres:5432/rag_assistant_dev"
        ),
        "APP_ENV": "test",
        "AUTH_ENABLED": "true",
        "AUTH_ISSUER": ISSUER,
        "AUTH_AUDIENCE": AUDIENCE,
        "AUTH_JWKS_URL": f"{ISSUER}/jwks",
        "AUTH_ALLOW_INSECURE_HTTP": "true",
        "AUTH_JWKS_CACHE_SECONDS": "300",
        "EMBEDDING_PROVIDER": "openai",
        "RERANKER_PROVIDER": "local",
        "DOWNLOAD_RERANKER": "true",
    }


def compose(*arguments: str, capture: bool = False) -> str:
    completed = subprocess.run(
        ["docker", "compose", *arguments],
        cwd=REPOSITORY,
        env=environment(),
        check=True,
        capture_output=capture,
        text=True,
    )
    return completed.stdout if capture else ""


def start_stack() -> None:
    compose("up", "-d", "postgres", "redis")
    compose(
        "run",
        "--rm",
        "api",
        "uv",
        "run",
        "--no-sync",
        "alembic",
        "upgrade",
        "head",
    )
    compose("up", "--build", "-d", "api", "worker")


def wait_for_api() -> None:
    with httpx.Client(timeout=2) as client:
        for _ in range(120):
            try:
                if client.get(f"{API}/health/ready").status_code == 200:
                    return
            except httpx.HTTPError:
                pass
            time.sleep(0.5)
    raise RuntimeError("API readiness timed out")


def require(response: httpx.Response, expected: int) -> dict[str, object]:
    if response.status_code != expected:
        raise RuntimeError(
            f"unexpected_api_status={response.status_code} expected={expected}"
        )
    return response.json() if response.content else {}


async def seed() -> dict[str, object]:
    from app.bootstrap_identity import bind_identity
    from app.db.models import Collection, MembershipRole, Organization
    from app.db.session import session_factory

    tenant_a, tenant_b = uuid4(), uuid4()
    collection_a, collection_b = uuid4(), uuid4()
    subjects = {role: f"m12-{role}-{uuid4()}" for role in ("viewer", "editor", "admin")}
    async with session_factory() as session, session.begin():
        session.add_all(
            [
                Organization(
                    id=tenant_a, name="M12 tenant A", slug=f"m12-a-{tenant_a}"
                ),
                Organization(
                    id=tenant_b, name="M12 tenant B", slug=f"m12-b-{tenant_b}"
                ),
            ]
        )
        await session.flush()
        session.add_all(
            [
                Collection(id=collection_a, tenant_id=tenant_a, name="Lifecycle A"),
                Collection(id=collection_b, tenant_id=tenant_b, name="Lifecycle B"),
            ]
        )
    users = {}
    for role, subject in subjects.items():
        user_id, _ = await bind_identity(
            ISSUER, subject, tenant_a, MembershipRole(role)
        )
        users[role] = user_id
    return {
        "tenant_a": tenant_a,
        "tenant_b": tenant_b,
        "collection_a": collection_a,
        "collection_b": collection_b,
        "subjects": subjects,
        "users": users,
    }


async def snapshot(document_id: UUID) -> dict[str, object]:
    from app.db.models import (
        Document,
        DocumentChunk,
        DocumentIndexGeneration,
        DocumentSourceUnit,
        DocumentVersion,
        ProcessingJob,
    )
    from app.db.session import session_factory

    async with session_factory() as session:
        document = await session.get(Document, document_id)
        versions = (
            await session.scalars(
                select(DocumentVersion)
                .where(DocumentVersion.document_id == document_id)
                .order_by(DocumentVersion.version_number)
            )
        ).all()
        generations = (
            await session.scalars(
                select(DocumentIndexGeneration)
                .where(DocumentIndexGeneration.document_id == document_id)
                .order_by(
                    DocumentIndexGeneration.document_version_id,
                    DocumentIndexGeneration.generation_number,
                )
            )
        ).all()
        jobs = (
            await session.scalars(
                select(ProcessingJob)
                .where(ProcessingJob.document_id == document_id)
                .order_by(ProcessingJob.created_at)
            )
        ).all()
        return {
            "document_status": document.status.value if document else None,
            "active_version_id": str(document.active_version_id)
            if document and document.active_version_id
            else None,
            "legacy_private_fields_clear": bool(
                document
                and document.storage_key is None
                and document.filename is None
                and document.checksum_sha256 is None
                and document.document_metadata == {}
            ),
            "versions": [
                {
                    "id": str(item.id),
                    "number": item.version_number,
                    "status": item.status.value,
                    "active_generation_id": str(item.active_generation_id)
                    if item.active_generation_id
                    else None,
                }
                for item in versions
            ],
            "generations": [
                {
                    "id": str(item.id),
                    "version_id": str(item.document_version_id),
                    "number": item.generation_number,
                    "status": item.status.value,
                    "fingerprint": item.configuration_fingerprint,
                }
                for item in generations
            ],
            "jobs": [
                {
                    "operation": item.operation,
                    "status": item.status.value,
                    "attempts": item.attempt_count,
                    "failure_category": item.failure_category,
                }
                for item in jobs
            ],
            "source_units": int(
                await session.scalar(
                    select(func.count(DocumentSourceUnit.id)).where(
                        DocumentSourceUnit.document_id == document_id
                    )
                )
                or 0
            ),
            "chunks": int(
                await session.scalar(
                    select(func.count(DocumentChunk.id)).where(
                        DocumentChunk.document_id == document_id
                    )
                )
                or 0
            ),
            "embeddings": int(
                await session.scalar(
                    select(func.count(DocumentChunk.id)).where(
                        DocumentChunk.document_id == document_id,
                        DocumentChunk.embedding.is_not(None),
                    )
                )
                or 0
            ),
            "full_text": int(
                await session.scalar(
                    select(func.count(DocumentChunk.id)).where(
                        DocumentChunk.document_id == document_id,
                        DocumentChunk.search_vector.is_not(None),
                    )
                )
                or 0
            ),
        }


async def chunk_lifecycle(chunk_id: UUID) -> tuple[str, str]:
    from app.db.models import DocumentChunk
    from app.db.session import session_factory

    async with session_factory() as session:
        chunk = await session.get(DocumentChunk, chunk_id)
        if chunk is None:
            raise RuntimeError("retrieved chunk disappeared")
        return str(chunk.document_version_id), str(chunk.generation_id)


async def citation_state(message_id: UUID) -> dict[str, object]:
    from app.db.models import ConversationCitation
    from app.db.session import session_factory

    async with session_factory() as session:
        citation = await session.scalar(
            select(ConversationCitation).where(
                ConversationCitation.assistant_message_id == message_id
            )
        )
        if citation is None:
            raise RuntimeError("historical citation missing")
        return {
            "source_status": citation.source_status,
            "version_id": str(citation.document_version_id)
            if citation.document_version_id
            else None,
            "source_links_cleared": all(
                value is None
                for value in (
                    citation.tenant_id,
                    citation.document_id,
                    citation.document_version_id,
                    citation.generation_id,
                    citation.chunk_id,
                )
            ),
            "content_redacted": citation.exact_excerpt is None
            and citation.document_metadata == {},
        }


async def create_failed_candidate(state: dict[str, object], document_id: UUID) -> UUID:
    from app.core.config import get_settings
    from app.db.models import (
        Document,
        DocumentIndexGeneration,
        DocumentVersion,
        ProcessingJob,
    )
    from app.db.session import session_factory
    from app.lifecycle import index_configuration
    from app.storage import version_storage_key

    version_id, generation_id, job_id = uuid4(), uuid4(), uuid4()
    async with session_factory() as session, session.begin():
        document = await session.scalar(
            select(Document).where(Document.id == document_id).with_for_update()
        )
        if document is None:
            raise RuntimeError("document missing for failed-candidate simulation")
        number = document.next_version_number
        document.next_version_number += 1
        config = index_configuration("application/pdf", get_settings())
        session.add(
            DocumentVersion(
                id=version_id,
                tenant_id=state["tenant_a"],
                collection_id=state["collection_a"],
                document_id=document_id,
                version_number=number,
                storage_key=version_storage_key(
                    state["tenant_a"], document_id, version_id, ".pdf"
                ),
                checksum_sha256="f" * 64,
                filename="failed-candidate.pdf",
                content_type="application/pdf",
                size_bytes=1,
                requested_by_user_id=state["users"]["editor"],
            )
        )
        await session.flush()
        session.add(
            DocumentIndexGeneration(
                id=generation_id,
                tenant_id=state["tenant_a"],
                document_id=document_id,
                document_version_id=version_id,
                generation_number=1,
                processing_job_id=job_id,
                requested_by_user_id=state["users"]["editor"],
                **config.__dict__,
                configuration_fingerprint=config.fingerprint,
            )
        )
        await session.flush()
        session.add(
            ProcessingJob(
                id=job_id,
                tenant_id=state["tenant_a"],
                document_id=document_id,
                document_version_id=version_id,
                generation_id=generation_id,
                operation="replacement_ingestion",
                idempotency_key=f"failed-{job_id}",
                request_fingerprint="f" * 64,
                requested_by_user_id=state["users"]["editor"],
            )
        )
    compose(
        "exec",
        "-T",
        "worker",
        "uv",
        "run",
        "--no-sync",
        "python",
        "-c",
        (
            "from app.tasks import verify_original_task;"
            f"verify_original_task.apply_async(args=['{state['tenant_a']}',"
            f"'{document_id}','{job_id}'])"
        ),
    )
    return job_id


async def wait_job(
    client: httpx.AsyncClient, headers: dict[str, str], job_id: str
) -> dict[str, object]:
    for _ in range(240):
        body = require(
            await client.get(f"/processing-jobs/{job_id}", headers=headers), 200
        )
        if body["status"] in {"succeeded", "failed"}:
            return body
        await asyncio.sleep(0.25)
    raise RuntimeError("processing job timed out")


def make_docx(path: Path, marker: str, revision: str) -> None:
    document = DocxDocument()
    document.add_heading("Synthetic lifecycle policy", level=1)
    document.add_paragraph(f"{marker} applies to the {revision} synthetic policy.")
    document.add_paragraph("The synthetic review interval is thirty days.")
    document.save(path)


async def verify(private_key: object) -> dict[str, object]:
    state = await seed()
    tokens = {
        role: access_token(private_key, state["subjects"][role])
        for role in ("viewer", "editor", "admin")
    }
    headers = {
        role: {"Authorization": f"Bearer {value}"} for role, value in tokens.items()
    }
    collection = state["collection_a"]
    metrics: dict[str, object] = {"answer_calls": []}
    temporary = tempfile.TemporaryDirectory()
    directory = temporary.name
    async with httpx.AsyncClient(base_url=API, timeout=90) as client:
        root = Path(directory)
        version_one_file = root / "version-one.docx"
        version_two_file = root / "version-two.docx"
        make_docx(version_one_file, "LIFECYCLEALPHA17", "first")
        make_docx(version_two_file, "LIFECYCLEBETA29", "second")

        with version_one_file.open("rb") as upload:
            uploaded = require(
                await client.post(
                    f"/collections/{collection}/documents",
                    headers=headers["editor"],
                    files={
                        "file": (
                            version_one_file.name,
                            upload,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                ),
                202,
            )
        document_id = UUID(uploaded["document_id"])
        initial_job = await wait_job(client, headers["editor"], uploaded["job_id"])
        assert initial_job["status"] == "succeeded"
        initial = await snapshot(document_id)
        version_one = initial["active_version_id"]
        generation_one = initial["versions"][0]["active_generation_id"]

        query_one = {"query": "LIFECYCLEALPHA17", "top_k": 5}
        keyword_one = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_one,
            ),
            200,
        )
        semantic_one = require(
            await client.post(
                f"/collections/{collection}/semantic-search",
                headers=headers["viewer"],
                json=query_one,
            ),
            200,
        )
        hybrid_one = require(
            await client.post(
                f"/collections/{collection}/hybrid-search",
                headers=headers["viewer"],
                json=query_one,
            ),
            200,
        )
        assert (
            keyword_one["results"] and semantic_one["results"] and hybrid_one["results"]
        )
        answer_one = require(
            await client.post(
                f"/collections/{collection}/ask",
                headers=headers["viewer"],
                json={
                    "query": "What synthetic policy applies to LIFECYCLEALPHA17?",
                    "retrieval_count": 5,
                },
            ),
            200,
        )
        assert answer_one["status"] == "answered" and answer_one["citations"]
        assert answer_one["citations"][0]["document_version_id"] == version_one
        metrics["answer_calls"].append(
            {
                "model": answer_one["usage"]["actual_model"],
                "usage": answer_one["usage"],
                "latency_ms": answer_one["latency"]["generation_ms"],
            }
        )

        conversation = require(
            await client.post(
                f"/collections/{collection}/conversations", headers=headers["viewer"]
            ),
            201,
        )
        conversation_turn = require(
            await client.post(
                f"/collections/{collection}/conversations/{conversation['id']}/messages",
                headers={**headers["viewer"], "Idempotency-Key": "m12-history-v1"},
                json={
                    "query": "Which synthetic policy uses LIFECYCLEALPHA17?",
                    "top_k": 5,
                },
            ),
            200,
        )
        historical_message_id = UUID(conversation_turn["assistant_message_id"])
        assert (
            conversation_turn["answer"]["citations"][0]["document_version_id"]
            == version_one
        )
        metrics["answer_calls"].append(
            {
                "model": conversation_turn["answer"]["usage"]["actual_model"],
                "usage": conversation_turn["answer"]["usage"],
                "latency_ms": conversation_turn["answer"]["latency"]["generation_ms"],
            }
        )

        viewer_upload = await client.post(
            f"/collections/{collection}/documents/{document_id}/versions",
            headers={**headers["viewer"], "Idempotency-Key": "viewer-denied"},
            files={
                "file": (
                    version_two_file.name,
                    version_two_file.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        viewer_reindex = await client.post(
            f"/collections/{collection}/documents/{document_id}/reindex",
            headers={**headers["viewer"], "Idempotency-Key": "viewer-reindex"},
        )
        viewer_delete = await client.delete(
            f"/collections/{collection}/documents/{document_id}",
            headers=headers["viewer"],
        )
        editor_delete = await client.delete(
            f"/collections/{collection}/documents/{document_id}",
            headers=headers["editor"],
        )
        cross_tenant = await client.get(
            f"/collections/{state['collection_b']}/documents/{document_id}",
            headers=headers["admin"],
        )
        assert (
            viewer_upload.status_code,
            viewer_reindex.status_code,
            viewer_delete.status_code,
            editor_delete.status_code,
            cross_tenant.status_code,
        ) == (403, 403, 403, 403, 404)

        compose("stop", "worker")
        with version_two_file.open("rb") as upload:
            replacement = require(
                await client.post(
                    f"/collections/{collection}/documents/{document_id}/versions",
                    headers={**headers["editor"], "Idempotency-Key": "replace-v2"},
                    files={
                        "file": (
                            version_two_file.name,
                            upload,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                ),
                202,
            )
        during_replace = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_one,
            ),
            200,
        )
        assert during_replace["results"]
        assert (await snapshot(document_id))["active_version_id"] == version_one
        compose("start", "worker")
        replacement_job = await wait_job(
            client, headers["editor"], replacement["job_id"]
        )
        assert replacement_job["status"] == "succeeded"
        replaced = await snapshot(document_id)
        version_two = replaced["active_version_id"]
        assert version_two != version_one
        generation_two = next(
            item["active_generation_id"]
            for item in replaced["versions"]
            if item["id"] == version_two
        )
        with version_two_file.open("rb") as upload:
            replacement_replay = require(
                await client.post(
                    f"/collections/{collection}/documents/{document_id}/versions",
                    headers={**headers["editor"], "Idempotency-Key": "replace-v2"},
                    files={
                        "file": (
                            version_two_file.name,
                            upload,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                ),
                202,
            )
        assert replacement_replay["job_id"] == replacement["job_id"]

        old_absent = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_one,
            ),
            200,
        )
        query_two = {"query": "LIFECYCLEBETA29", "top_k": 5}
        keyword_two = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        semantic_two = require(
            await client.post(
                f"/collections/{collection}/semantic-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        hybrid_two = require(
            await client.post(
                f"/collections/{collection}/hybrid-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        assert (
            not old_absent["results"]
            and keyword_two["results"]
            and semantic_two["results"]
            and hybrid_two["results"]
        )
        answer_two = require(
            await client.post(
                f"/collections/{collection}/ask",
                headers=headers["viewer"],
                json={
                    "query": "What synthetic policy applies to LIFECYCLEBETA29?",
                    "retrieval_count": 5,
                },
            ),
            200,
        )
        assert (
            answer_two["status"] == "answered"
            and answer_two["citations"][0]["document_version_id"] == version_two
        )
        metrics["answer_calls"].append(
            {
                "model": answer_two["usage"]["actual_model"],
                "usage": answer_two["usage"],
                "latency_ms": answer_two["latency"]["generation_ms"],
            }
        )
        historical_before_delete = await citation_state(historical_message_id)
        assert historical_before_delete["version_id"] == version_one

        compose("stop", "worker")
        reindex = require(
            await client.post(
                f"/collections/{collection}/documents/{document_id}/reindex",
                headers={**headers["editor"], "Idempotency-Key": "reindex-v2"},
            ),
            202,
        )
        during_reindex = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        assert during_reindex["results"]
        compose("start", "worker")
        reindex_job = await wait_job(client, headers["editor"], reindex["job_id"])
        assert reindex_job["status"] == "succeeded"
        reindexed = await snapshot(document_id)
        active_generation_after_reindex = next(
            item["active_generation_id"]
            for item in reindexed["versions"]
            if item["id"] == version_two
        )
        assert (
            reindexed["active_version_id"] == version_two
            and active_generation_after_reindex != generation_two
        )
        reindex_replay = require(
            await client.post(
                f"/collections/{collection}/documents/{document_id}/reindex",
                headers={**headers["editor"], "Idempotency-Key": "reindex-v2"},
            ),
            202,
        )
        assert reindex_replay["job_id"] == reindex["job_id"]
        hybrid_after_reindex = require(
            await client.post(
                f"/collections/{collection}/hybrid-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        lifecycle_rows = [
            await chunk_lifecycle(UUID(row["chunk_id"]))
            for row in hybrid_after_reindex["results"]
        ]
        assert (
            lifecycle_rows
            and {row[0] for row in lifecycle_rows} == {version_two}
            and {row[1] for row in lifecycle_rows} == {active_generation_after_reindex}
        )
        assert len({row["chunk_id"] for row in hybrid_after_reindex["results"]}) == len(
            hybrid_after_reindex["results"]
        )

        failed_job_id = await create_failed_candidate(state, document_id)
        failed_job = await wait_job(client, headers["editor"], str(failed_job_id))
        assert failed_job["status"] == "failed"
        after_failure = await snapshot(document_id)
        assert after_failure["active_version_id"] == version_two
        assert (
            next(
                item["active_generation_id"]
                for item in after_failure["versions"]
                if item["id"] == version_two
            )
            == active_generation_after_reindex
        )

        compose("stop", "worker")
        deletion = require(
            await client.delete(
                f"/collections/{collection}/documents/{document_id}",
                headers=headers["admin"],
            ),
            202,
        )
        immediate_keyword = require(
            await client.post(
                f"/collections/{collection}/keyword-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        immediate_semantic = require(
            await client.post(
                f"/collections/{collection}/semantic-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        immediate_hybrid = require(
            await client.post(
                f"/collections/{collection}/hybrid-search",
                headers=headers["viewer"],
                json=query_two,
            ),
            200,
        )
        immediate_ask = require(
            await client.post(
                f"/collections/{collection}/ask",
                headers=headers["viewer"],
                json={
                    "query": "What synthetic policy applies to LIFECYCLEBETA29?",
                    "retrieval_count": 5,
                },
            ),
            200,
        )
        empty_conversation = require(
            await client.post(
                f"/collections/{collection}/conversations", headers=headers["viewer"]
            ),
            201,
        )
        immediate_turn = require(
            await client.post(
                f"/collections/{collection}/conversations/{empty_conversation['id']}/messages",
                headers={**headers["viewer"], "Idempotency-Key": "after-delete"},
                json={"query": "What synthetic policy applies?", "top_k": 5},
            ),
            200,
        )
        assert (
            not immediate_keyword["results"]
            and not immediate_semantic["results"]
            and not immediate_hybrid["results"]
        )
        assert immediate_ask["status"] == "insufficient_context"
        assert immediate_turn["answer"]["status"] == "insufficient_context"
        compose("start", "worker")
        deletion_job = await wait_job(client, headers["admin"], deletion["job_id"])
        assert deletion_job["status"] == "succeeded"
        deletion_replay = require(
            await client.delete(
                f"/collections/{collection}/documents/{document_id}",
                headers=headers["admin"],
            ),
            202,
        )
        assert deletion_replay["job_id"] == deletion["job_id"]
        deleted = await snapshot(document_id)
        citation_after_delete = await citation_state(historical_message_id)
        assert (
            deleted["document_status"] == "deleted"
            and deleted["legacy_private_fields_clear"]
        )
        assert not deleted["versions"] and not deleted["generations"]
        assert (
            deleted["source_units"]
            == deleted["chunks"]
            == deleted["embeddings"]
            == deleted["full_text"]
            == 0
        )
        assert (
            citation_after_delete["source_status"] == "deleted"
            and citation_after_delete["source_links_cleared"]
            and citation_after_delete["content_redacted"]
        )

        metrics.update(
            {
                "document_id": str(document_id),
                "version_one": version_one,
                "generation_one": generation_one,
                "version_two": version_two,
                "generation_two": generation_two,
                "reindex_generation": active_generation_after_reindex,
                "initial": initial,
                "replaced": replaced,
                "reindexed": reindexed,
                "after_failure": after_failure,
                "deleted": deleted,
                "permissions": (
                    "viewer_read_only=passed editor_lifecycle=passed "
                    "admin_delete=passed cross_tenant_404=passed"
                ),
                "idempotency": "replacement=passed reindex=passed deletion=passed",
                "historical_citation": (
                    "superseded_resolved=passed deletion_tombstone=passed"
                ),
            }
        )
    temporary.cleanup()
    return metrics


def provider_and_log_metrics() -> dict[str, object]:
    logs = compose("logs", "--no-color", "api", "worker", capture=True)
    embedding = re.findall(
        r"Embedding provider completed model=(\S+) inputs=(\d+) "
        r"input_tokens=(\d+) total_tokens=(\d+) latency_ms=([0-9.]+)",
        logs,
    )
    forbidden = {
        "authorization_header": bool(re.search(r"Authorization:\s*Bearer", logs, re.I)),
        "bearer_token": bool(re.search(r"Bearer\s+eyJ", logs)),
        "openai_key": bool(re.search(r"sk-[A-Za-z0-9_-]{12,}", logs)),
        "source_text": "LIFECYCLEALPHA17" in logs or "LIFECYCLEBETA29" in logs,
        "vector": bool(re.search(r"\[(?:-?\d+\.\d+,\s*){20,}", logs)),
        "storage_path": "/data/documents" in logs,
        "raw_provider_response": "output_parsed" in logs or "response.data" in logs,
    }
    if any(forbidden.values()):
        raise RuntimeError("redacted log inspection failed")
    return {
        "embedding_calls": len(embedding),
        "embedding_model": sorted({row[0] for row in embedding}),
        "embedding_inputs": sum(int(row[1]) for row in embedding),
        "embedding_input_tokens": sum(int(row[2]) for row in embedding),
        "embedding_total_tokens": sum(int(row[3]) for row in embedding),
        "embedding_latencies_ms": [float(row[4]) for row in embedding],
        "log_secret_scan": "passed",
    }


def container_health() -> str:
    expected = {"api", "worker", "postgres", "redis"}
    for _ in range(120):
        value = compose("ps", "--format", "json", capture=True)
        rows = [json.loads(line) for line in value.splitlines() if line.strip()]
        healthy = {
            row["Service"] for row in rows if row.get("Health") == "healthy"
        }
        if healthy == expected:
            return "api=healthy worker=healthy postgres=healthy redis=healthy"
        time.sleep(0.5)
    raise RuntimeError("container health check failed")


def main() -> None:
    os.environ["DATABASE_URL"] = (
        "postgresql+asyncpg://rag_assistant_dev:rag_assistant_dev@"
        f"127.0.0.1:{POSTGRES_PORT}/rag_assistant_dev"
    )
    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    QuietJWKSHandler.jwks = {"keys": [public_jwk(private_key)]}
    server = ThreadingHTTPServer(("0.0.0.0", 8765), QuietJWKSHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        start_stack()
        wait_for_api()
        metrics = asyncio.run(verify(private_key))
        reconciliation = compose(
            "exec",
            "-T",
            "worker",
            "uv",
            "run",
            "--no-sync",
            "python",
            "-m",
            "app.storage_reconciliation",
            capture=True,
        ).strip()
        if (
            "missing_objects=0" not in reconciliation
            or "orphan_objects=0" not in reconciliation
        ):
            raise RuntimeError("storage reconciliation failed")
        metrics["storage_reconciliation"] = reconciliation
        metrics["containers"] = container_health()
        metrics.update(provider_and_log_metrics())
        print(json.dumps(metrics, sort_keys=True, indent=2))
    finally:
        server.shutdown()
        server.server_close()
        compose("down", "-v", "--remove-orphans")


if __name__ == "__main__":
    main()
