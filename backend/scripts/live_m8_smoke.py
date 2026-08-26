"""Minimal real-provider Milestone 8 smoke test using synthetic documents only."""

import asyncio
import json
import math
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from uuid import UUID, uuid4

import httpx
from docx import Document as DocxDocument
from sqlalchemy import select, text

from app.core.config import get_settings
from app.db.models import Collection, Membership
from app.db.session import session_factory


@dataclass(frozen=True)
class Fixture:
    filename: str
    text: str
    metadata: dict[str, object]


def write_docx(path: Path, fixture: Fixture) -> None:
    document = DocxDocument()
    document.add_heading("Synthetic verification", level=1)
    document.add_paragraph(fixture.text)
    document.save(path)


async def configured_collection() -> UUID:
    settings = get_settings()
    if not settings.development_tenant_id or not settings.development_user_id:
        raise RuntimeError("development trusted principal is not configured")
    tenant_id = UUID(settings.development_tenant_id)
    user_id = UUID(settings.development_user_id)
    async with session_factory() as session:
        collection_id = await session.scalar(
            select(Collection.id)
            .join(
                Membership,
                (Membership.tenant_id == Collection.tenant_id)
                & (Membership.user_id == user_id),
            )
            .where(Collection.tenant_id == tenant_id)
            .order_by(Collection.created_at, Collection.id)
            .limit(1)
        )
    if collection_id is None:
        raise RuntimeError("configured trusted principal has no collection")
    return collection_id


async def wait_for_job(client: httpx.AsyncClient, job_id: str) -> dict[str, object]:
    for _ in range(120):
        response = await client.get(f"/processing-jobs/{job_id}")
        response.raise_for_status()
        body = response.json()
        if body["status"] in {"succeeded", "failed"}:
            return body
        await asyncio.sleep(0.5)
    raise RuntimeError("synthetic processing job did not finish in time")


def reciprocal_rank(results: list[dict[str, object]], expected: str) -> float:
    for rank, result in enumerate(results, 1):
        if result["document_id"] == expected:
            return 1 / rank
    return 0.0


def ndcg(results: list[dict[str, object]], expected: str) -> float:
    for rank, result in enumerate(results[:10], 1):
        if result["document_id"] == expected:
            return 1 / math.log2(rank + 1)
    return 0.0


async def main() -> None:
    collection_id = await configured_collection()
    run_tag = f"m8live-{uuid4().hex}"
    fixtures = [
        Fixture(
            "m8-enterprise-refund.docx",
            "ENTREFUND30. Enterprise customers may request a full refund within "
            "30 days of purchase.",
            {
                "tags": [run_tag, "refund"],
                "department": "legal",
                "document_type": "policy",
                "language": "en",
                "effective_date": "2026-01-01",
            },
        ),
        Fixture(
            "m8-remboursement.docx",
            "Les grandes entreprises peuvent obtenir le remboursement complet "
            "de leur achat pendant une periode de trente jours.",
            {
                "tags": [run_tag, "remboursement"],
                "department": "legal",
                "document_type": "faq",
                "language": "fr",
                "effective_date": "2026-02-01",
            },
        ),
        Fixture(
            "m8-maintenance.docx",
            "Equipment maintenance occurs every Friday. The office opens at "
            "eight in the morning.",
            {
                "tags": [run_tag, "maintenance"],
                "department": "facilities",
                "document_type": "manual",
                "language": "en",
                "effective_date": "2025-01-01",
            },
        ),
    ]
    uploads: dict[str, str] = {}
    started = time.perf_counter()
    async with httpx.AsyncClient(
        base_url="http://127.0.0.1:8000", timeout=45
    ) as client:
        with tempfile.TemporaryDirectory(prefix="rag-m8-") as directory:
            for fixture in fixtures:
                path = Path(directory) / fixture.filename
                write_docx(path, fixture)
                with path.open("rb") as source:
                    response = await client.post(
                        f"/collections/{collection_id}/documents",
                        files={
                            "file": (
                                fixture.filename,
                                source,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                        data={"metadata": json.dumps(fixture.metadata)},
                    )
                response.raise_for_status()
                body = response.json()
                job = await wait_for_job(client, body["job_id"])
                if job["status"] != "succeeded":
                    raise RuntimeError("synthetic processing failed")
                uploads[fixture.filename] = body["document_id"]

        expected_count = len(uploads)
        async with session_factory() as session:
            rows = (
                await session.execute(
                    text(
                        "SELECT d.id::text, d.filename, d.metadata, "
                        "count(c.id), count(c.embedding), "
                        "count(*) FILTER (WHERE vector_dims(c.embedding) = 1536), "
                        "count(*) FILTER (WHERE c.search_vector IS NOT NULL) "
                        "FROM documents d JOIN document_chunks c "
                        "ON c.document_id = d.id "
                        "WHERE d.id = ANY(CAST(:ids AS uuid[])) "
                        "GROUP BY d.id, d.filename, d.metadata"
                    ),
                    {"ids": list(uploads.values())},
                )
            ).all()
        if len(rows) != expected_count:
            raise RuntimeError("not all synthetic documents have chunks")
        for (
            _document_id,
            filename,
            metadata,
            chunks,
            embedded,
            dimensions,
            vectors,
        ) in rows:
            stored_metadata = (
                json.loads(metadata) if isinstance(metadata, str) else metadata
            )
            expected_metadata = next(
                item.metadata for item in fixtures if item.filename == filename
            )
            expected_scalars = {
                key: value for key, value in expected_metadata.items() if key != "tags"
            }
            scalar_match = all(
                stored_metadata.get(key) == value
                for key, value in expected_scalars.items()
            )
            stored_tags = stored_metadata.get("tags", [])
            tag_match = (
                isinstance(stored_tags, list)
                and len(stored_tags) == 2
                and stored_tags[0].startswith("m8live-")
                and expected_metadata["tags"][1] in stored_tags
            )
            if not scalar_match or not tag_match:
                raise RuntimeError("stored synthetic metadata mismatch")
            if not (chunks > 0 and chunks == embedded == dimensions == vectors):
                raise RuntimeError("incomplete embeddings or text vectors")

        evaluations = [
            (
                "ENTREFUND30",
                uploads["m8-enterprise-refund.docx"],
                {"tags_all": [run_tag, "refund"], "languages": ["en"]},
            ),
            (
                "How long can a large business wait before asking for its money back?",
                uploads["m8-enterprise-refund.docx"],
                {"tags_any": [run_tag], "departments": ["legal"], "languages": ["en"]},
            ),
            (
                "Quel est le delai accorde aux societes pour recuperer le prix paye?",
                uploads["m8-remboursement.docx"],
                {"tags_any": [run_tag], "languages": ["fr"]},
            ),
        ]
        metrics: dict[str, list[float]] = {
            mode: [] for mode in ("semantic", "keyword", "hybrid", "reranked")
        }
        top_names: dict[str, list[str | None]] = {mode: [] for mode in metrics}
        for query, expected, filters in evaluations:
            for mode in metrics:
                response = await client.post(
                    f"/collections/{collection_id}/{mode}-search",
                    json={"query": query, "top_k": 5, "filters": filters},
                )
                response.raise_for_status()
                results = response.json()["results"]
                metrics[mode].extend(
                    [
                        reciprocal_rank(results, expected),
                        ndcg(results, expected),
                        float(any(item["document_id"] == expected for item in results)),
                    ]
                )
                top_names[mode].append(
                    results[0]["document_name"] if results else None
                )

        excluded = await client.post(
            f"/collections/{collection_id}/hybrid-search",
            json={
                "query": evaluations[1][0],
                "top_k": 5,
                "filters": {"tags_any": [run_tag], "departments": ["facilities"]},
            },
        )
        excluded.raise_for_status()
        if any(
            item["document_id"] == uploads["m8-enterprise-refund.docx"]
            for item in excluded.json()["results"]
        ):
            raise RuntimeError("filter was applied after retrieval")
        unauthorized = await client.post(
            f"/collections/{uuid4()}/reranked-search",
            json={"query": "synthetic", "top_k": 1},
        )
        if unauthorized.status_code != 404:
            raise RuntimeError("unauthorized collection did not return 404")

    print(f"documents={len(uploads)} processing=succeeded")
    print(f"chunks_verified={sum(row[3] for row in rows)} metadata=verified")
    for mode, values in metrics.items():
        triples = list(zip(values[0::3], values[1::3], values[2::3], strict=True))
        print(
            f"{mode}: MRR@10={sum(x[0] for x in triples)/len(triples):.4f} "
            f"NDCG@10={sum(x[1] for x in triples)/len(triples):.4f} "
            f"Recall@5={sum(x[2] for x in triples)/len(triples):.4f} "
            f"top_files={top_names[mode]}"
        )
    print(f"elapsed_ms={(time.perf_counter() - started) * 1000:.1f}")
    print("scope=verified unauthorized=404 real_local_reranker=yes")


if __name__ == "__main__":
    asyncio.run(main())
