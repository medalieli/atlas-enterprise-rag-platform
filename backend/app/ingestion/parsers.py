import time
from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path

from docx import Document as OpenXmlDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.ingestion.types import ParsedDocument, SourceBlock, SourceLocation, SourceUnit

PDF_PARSER_VERSION = "pypdf-6-v1"
DOCX_PARSER_VERSION = "python-docx-1-v1"


class PermanentParserError(Exception):
    pass


class ParserLimitError(PermanentParserError):
    pass


class OcrRequiredError(PermanentParserError):
    pass


@dataclass(frozen=True)
class ParserLimits:
    max_pdf_pages: int
    max_extracted_chars: int
    max_pdf_stream_bytes: int
    max_seconds: int


def _check_time(started: float, limits: ParserLimits) -> None:
    if time.monotonic() - started > limits.max_seconds:
        raise ParserLimitError("Document parsing exceeded the configured time limit")


def _encoded_stream_bytes(page: object) -> int:
    raw = page.raw_get("/Contents")  # type: ignore[attr-defined]
    references = raw if isinstance(raw, list) else [raw]
    total = 0
    for reference in references:
        stream = reference.get_object()
        data = getattr(stream, "_data", None)
        if isinstance(data, bytes):
            total += len(data)
    return total


def parse_pdf(path: Path, limits: ParserLimits) -> ParsedDocument:
    started = time.monotonic()
    try:
        reader = PdfReader(path, strict=True)
    except (PdfReadError, OSError, ValueError) as exc:
        raise PermanentParserError("PDF is corrupt or malformed") from exc
    if reader.is_encrypted:
        raise PermanentParserError("Encrypted PDFs are not supported")
    if len(reader.pages) > limits.max_pdf_pages:
        raise ParserLimitError("PDF exceeds the configured page limit")
    units: list[SourceUnit] = []
    total_chars = 0
    total_stream_bytes = 0
    try:
        for page_index, page in enumerate(reader.pages):
            _check_time(started, limits)
            if "/Contents" in page:
                total_stream_bytes += _encoded_stream_bytes(page)
                if total_stream_bytes > limits.max_pdf_stream_bytes:
                    raise ParserLimitError(
                        "PDF content streams exceed the configured limit"
                    )
            text = (
                page.extract_text(extraction_mode="layout") or ""
                if "/Contents" in page
                else ""
            )
            total_chars += len(text)
            if total_chars > limits.max_extracted_chars:
                raise ParserLimitError("Extracted text exceeds the configured limit")
            units.append(
                SourceUnit(
                    unit_index=page_index,
                    location=SourceLocation("pdf", page_number=page_index + 1),
                    blocks=(SourceBlock(0, "page", text),),
                )
            )
    except ParserLimitError:
        raise
    except Exception as exc:
        raise PermanentParserError("PDF content could not be extracted") from exc
    if not any(unit.blocks[0].text.strip() for unit in units):
        raise OcrRequiredError("PDF has no extractable text; OCR is required")
    return ParsedDocument(PDF_PARSER_VERSION, tuple(units))


def _table_text(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _docx_blocks(document: OpenXmlDocument) -> Iterator[tuple[str, str, str | None]]:
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            style = item.style.name if item.style is not None else ""
            kind = "heading" if style.lower().startswith("heading") else "paragraph"
            if (
                item._p.pPr is not None and item._p.pPr.numPr is not None
            ) or style.lower().startswith("list"):
                kind = "list"
                yield kind, f"- {item.text}", style
            else:
                yield kind, item.text, style
        elif isinstance(item, Table):
            yield "table", _table_text(item), None


def parse_docx(path: Path, limits: ParserLimits) -> ParsedDocument:
    started = time.monotonic()
    try:
        document = OpenXmlDocument(path)
    except Exception as exc:
        raise PermanentParserError("DOCX is corrupt or malformed") from exc
    units: list[SourceUnit] = []
    section_path: list[str] = []
    blocks: list[SourceBlock] = []
    chars = 0

    def flush() -> None:
        if blocks:
            units.append(
                SourceUnit(
                    len(units),
                    SourceLocation("docx", section_path=tuple(section_path)),
                    tuple(blocks),
                )
            )
            blocks.clear()

    for kind, text, style in _docx_blocks(document):
        _check_time(started, limits)
        chars += len(text)
        if chars > limits.max_extracted_chars:
            raise ParserLimitError("Extracted text exceeds the configured limit")
        if kind == "heading" and text.strip():
            flush()
            try:
                level = int((style or "Heading 1").rsplit(" ", 1)[-1])
            except ValueError:
                level = 1
            section_path[:] = section_path[: max(0, level - 1)] + [text.strip()]
        blocks.append(SourceBlock(len(blocks), kind, text))
    flush()
    if not units or not any(
        block.text.strip() for unit in units for block in unit.blocks
    ):
        raise PermanentParserError("DOCX has no extractable text")
    return ParsedDocument(DOCX_PARSER_VERSION, tuple(units))


def parse_document(
    path: Path, content_type: str, limits: ParserLimits
) -> ParsedDocument:
    if content_type == "application/pdf":
        return parse_pdf(path, limits)
    if (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return parse_docx(path, limits)
    raise PermanentParserError("Unsupported parser content type")
