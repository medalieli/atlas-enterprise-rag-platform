import io

from docx import Document


def pdf_bytes(pages: list[str]) -> bytes:
    """Create a small deterministic text PDF without a fixture-generation dependency."""
    objects: list[bytes] = []
    page_ids = [3 + index * 2 for index in range(len(pages))]
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    for index, text in enumerate(pages):
        page_id = page_ids[index]
        content_id = page_id + 1
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
                f"/BaseFont /Helvetica >> >> >> /Contents {content_id} 0 R >>"
            ).encode()
        )
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_id, value in enumerate(objects, 1):
        offsets.append(len(output))
        output.extend(f"{object_id} 0 obj\n".encode() + value + b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref}\n%%EOF\n"
    )
    output.extend(trailer.encode())
    return bytes(output)


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Policy", level=1)
    document.add_paragraph("First paragraph.")
    document.add_paragraph("First item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.add_heading("Details", level=2)
    document.add_paragraph("Second section.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
